#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学生手册Word文档提取工具 - 干净版本
设计师：内师智能体系统 (￣▽￣)ﾉ
优势：使用Word文档提取，避免PDF提取的噪音问题
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
import json
from docx import Document

def extract_from_word(docx_path, output_path):
    """从Word文档提取完整内容"""
    print(f"正在读取Word文档：{docx_path}")

    try:
        doc = Document(docx_path)
        print(f"✓ 文档打开成功")

        # 提取所有段落
        content = {
            "total_pages": len(doc.paragraphs),  # Word文档按段落计算"页数"
            "pages": []
        }

        current_page_num = 1
        current_page_text = []
        char_count = 0
        CHARS_PER_PAGE = 1500  # 每页大约1500字符
        PARAS_PER_PAGE = 10     # 每页最多10个段落作为兜底

        print(f"\n✓ 正在提取内容...")
        print(f"✓ 总段落数：{len(doc.paragraphs)}")

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_page_text.append(text)
                char_count += len(text)

                # 当达到每页字符数或段落数时，创建新页
                if char_count >= CHARS_PER_PAGE or len(current_page_text) >= PARAS_PER_PAGE:
                    content["pages"].append({
                        "page_num": current_page_num,
                        "text": "\n".join(current_page_text)
                    })

                    current_page_num += 1
                    current_page_text = []
                    char_count = 0

        # 添加最后一页
        if current_page_text:
            content["pages"].append({
                "page_num": current_page_num,
                "text": "\n".join(current_page_text)
            })

        # 保存为JSON
        print(f"\n✓ 正在保存到文件...")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)

        print(f"✓ 提取完成！")
        print(f"  - 总段落数：{len(doc.paragraphs)}")
        print(f"  - 总页数：{content['total_pages']}")
        print(f"  - 提取页数：{len(content['pages'])}")
        print(f"  - 保存路径：{output_path}")

        return content

    except Exception as e:
        print(f"❌ 提取失败：{e}")
        import traceback
        traceback.print_exc()
        return None

def analyze_quality(content):
    """分析提取内容的质量"""
    print(f"\n✓ 正在分析内容质量...")

    total_chars = 0
    noise_lines = 0
    punctuation_lines = 0

    for page in content["pages"]:
        text = page["text"]
        total_chars += len(text)

        lines = text.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 检查是否为纯标点符号行（噪音）
            if len(stripped) <= 10 and all(c in '。，、；：,.;:!?！？' for c in stripped):
                noise_lines += 1

            # 检查是否包含标点符号
            if any(c in '。，、；：,.;:!?！？' for c in stripped):
                punctuation_lines += 1

    print(f"\n📊 质量统计：")
    print(f"  - 总字符数：{total_chars}")
    print(f"  - 总行数：{sum(len(p['text'].split('\\n')) for p in content['pages'])}")
    print(f"  - 噪音行数：{noise_lines}")
    print(f"  - 包含标点的行数：{punctuation_lines}")

    if noise_lines == 0:
        print(f"\n🎉 优秀！没有发现噪音行！文本质量很高！(￣▽￣)ﾉ")
    else:
        print(f"\n⚠️ 发现 {noise_lines} 个噪音行，可能需要进一步清理")

if __name__ == "__main__":
    print("=" * 60)
    print("🚀 学生手册Word文档提取工具")
    print("设计师：内师智能体系统 (￣▽￣)ﾉ")
    print("=" * 60)

    # 从命令行参数获取路径
    if len(sys.argv) >= 3:
        docx_path = sys.argv[1]
        output_path = sys.argv[2]
    else:
        print("用法: python extract_student_handbook_word.py <docx_path> <output_path>")
        print("示例: python extract_student_handbook_word.py input.docx output.json")
        sys.exit(1)

    print(f"输入文件: {docx_path}")
    print(f"输出文件: {output_path}")

    # 提取内容
    content = extract_from_word(docx_path, output_path)

    if content:
        # 分析质量
        analyze_quality(content)

        print(f"\n" + "=" * 60)
        print(f"✅ 提取完成！干净的JSON文件已保存：")
        print(f"   {output_path}")
        print(f"=" * 60)
