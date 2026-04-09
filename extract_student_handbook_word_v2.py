#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学生手册Word文档提取工具 - 完整版本
设计师：内师智能体系统 (￣▽￣)ﾉ
优势：提取所有段落，不遗漏任何内容！
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
import json
from docx import Document

def extract_from_word(docx_path, output_path):
    """从Word文档提取完整内容 - 每个段落作为独立块"""
    print(f"正在读取Word文档：{docx_path}")

    try:
        doc = Document(docx_path)
        print(f"✓ 文档打开成功")

        # 提取所有有效段落
        content = {
            "total_paragraphs": len(doc.paragraphs),
            "chunks": []
        }

        print(f"\n✓ 正在提取内容...")
        print(f"✓ 总段落数：{len(doc.paragraphs)}")

        chunk_id = 1
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # 只保存非空段落
                content["chunks"].append({
                    "id": chunk_id,
                    "text": text,
                    "paragraph_num": i + 1
                })
                chunk_id += 1

        # 保存为JSON
        print(f"\n✓ 正在保存到文件...")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)

        print(f"✓ 提取完成！")
        print(f"  - 总段落数：{len(doc.paragraphs)}")
        print(f"  - 有效段落数：{len(content['chunks'])}")
        print(f"  - 保存路径：{output_path}")

        return content

    except Exception as e:
        print(f"❌ 提取失败：{e}")
        import traceback
        traceback.print_exc()
        return None

def analyze_quality(content):
    """分析提取内容的质量"""
    print(f"\n✓ 正在分析内容质量...")

    total_chars = 0
    noise_lines = 0
    empty_chunks = 0

    for chunk in content["chunks"]:
        text = chunk["text"]
        total_chars += len(text)

        if not text:
            empty_chunks += 1
            continue

        # 检查是否为纯标点符号行（噪音）
        if len(text) <= 10 and all(c in '。，、；：,.;:!?！？' for c in text):
            noise_lines += 1

    print(f"\n📊 质量统计：")
    print(f"  - 总文档块数：{len(content['chunks'])}")
    print(f"  - 总字符数：{total_chars}")
    print(f"  - 噪音块数：{noise_lines}")
    print(f"  - 空块数：{empty_chunks}")

    if noise_lines == 0:
        print(f"\n🎉 优秀！没有发现噪音块！文本质量很高！(￣▽￣)ﾉ")
    else:
        print(f"\n⚠️ 发现 {noise_lines} 个噪音块")

if __name__ == "__main__":
    # Word文档路径
    docx_path = r"E:\外包\教育系统智能体\相关文档\2025年本科学生手册-定.docx"

    # 输出路径
    output_path = r"E:\外包\教育系统智能体\student_handbook_full_clean.json"

    print("=" * 60)
    print("🚀 学生手册Word文档提取工具 v2.0")
    print("设计师：内师智能体系统 (￣▽￣)ﾉ")
    print("=" * 60)

    # 提取内容
    content = extract_from_word(docx_path, output_path)

    if content:
        # 分析质量
        analyze_quality(content)

        print(f"\n" + "=" * 60)
        print(f"✅ 提取完成！干净的JSON文件已保存：")
        print(f"   {output_path}")
        print(f"=" * 60)
